from flask import Blueprint, request, jsonify
from ..db import get_conn
from ..services.retreival import build_faiss_index
from ..services.text import normalizar_idioma
from ..config import Config
import os
import traceback

# Importações opcionais para evitar erros se os pacotes não estiverem instalados
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

app = Blueprint('uploads', __name__)


@app.route("/upload-pdf", methods=["POST"])
def upload_pdf():
    if not PDF_AVAILABLE:
        return jsonify({"success": False, "error": "PyPDF2 não está instalado."}), 500
    
    conn = get_conn()
    cur = conn.cursor()
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "Ficheiro não enviado."}), 400
    files = request.files.getlist('file')
    chatbot_id = request.form.get("chatbot_id")
    if not chatbot_id:
        return jsonify({"success": False, "error": "Chatbot ID não fornecido."}), 400
    uploaded_pdf_ids = []
    try:
        if not os.path.exists(Config.PDF_STORAGE_PATH):
            os.makedirs(Config.PDF_STORAGE_PATH)
        for file in files:
            filename = file.filename
            if not filename.lower().endswith('.pdf'):
                return jsonify({"success": False, "error": "Apenas ficheiros PDF são permitidos."}), 400
            file_path = os.path.join(Config.PDF_STORAGE_PATH, filename)
            file.save(file_path)
            file.close()
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                if reader.is_encrypted:
                    return jsonify({"success": False, "error": f"O PDF '{filename}' está protegido por senha."}), 400
                if not reader.pages:
                    return jsonify({"success": False, "error": f"O PDF '{filename}' está vazio ou corrompido."}), 400
            cur.execute(
                "INSERT INTO PDF_Documents (chatbot_id, filename, file_path) VALUES (%s, %s, %s) RETURNING pdf_id",
                (chatbot_id, filename, file_path)
            )
            pdf_id = cur.fetchone()[0]
            uploaded_pdf_ids.append(pdf_id)
        conn.commit()
        return jsonify({"success": True, "pdf_ids": uploaded_pdf_ids, "message": "PDF(s) carregado(s) com sucesso."})
    except PyPDF2.errors.PdfReadError:
        return jsonify({"success": False, "error": "Erro ao ler o PDF. Verifique se o arquivo não está corrompido."}), 400
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        cur.close()
        conn.close()

@app.route("/upload-faq-docx", methods=["POST"])
def upload_faq_docx():
    if not DOCX_AVAILABLE:
        return jsonify({"success": False, "error": "python-docx não está instalado."}), 500
    
    conn = get_conn()
    cur = conn.cursor()
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "Ficheiro não enviado."}), 400
    file = request.files['file']
    chatbot_id_raw = request.form.get("chatbot_id")
    if not chatbot_id_raw:
        return jsonify({"success": False, "error": "Chatbot ID não fornecido."}), 400
    try:
        doc = docx.Document(file)
        dados = {}
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    chave_raw = row.cells[0].text.strip().lower().replace("\u2019", "'")
                    valor = row.cells[1].text.strip()
                    chave = chave_raw.replace(":", "").strip()
                    if chave and valor:
                        dados[chave] = valor
        if not dados.get("designação da faq") or not dados.get("questão") or not dados.get("resposta"):
            raise Exception("Faltam campos obrigatórios: designação, questão ou resposta.")
        designacao = dados.get("designação da faq")
        pergunta = dados.get("questão")
        resposta = dados.get("resposta")
        categoria = dados.get("categoria")
        idioma_lido = dados.get("idioma", "Português")
        idioma = normalizar_idioma(idioma_lido)
        links_documentos = ""
        for key in ["documentos associados", "links de documentos"]:
            if key in dados:
                links_documentos = dados[key]
                break
        chatbot_ids = []
        if chatbot_id_raw == "todos":
            cur.execute("SELECT chatbot_id FROM Chatbot")
            chatbot_ids = [row[0] for row in cur.fetchall()]
        else:
            chatbot_ids = [int(chatbot_id_raw)]
        for chatbot_id in chatbot_ids:
            cur.execute("""
                SELECT faq_id FROM FAQ
                WHERE chatbot_id = %s AND designacao = %s AND pergunta = %s AND resposta = %s AND idioma = %s
            """, (chatbot_id, designacao, pergunta, resposta, idioma))
            if cur.fetchone():
                continue
            cur.execute("""
                INSERT INTO FAQ (chatbot_id, designacao, pergunta, resposta, idioma, links_documentos)
                VALUES (%s, %s, %s, %s, %s, %s)
                RETURNING faq_id
            """, (chatbot_id, designacao, pergunta, resposta, idioma, links_documentos))
            faq_id = cur.fetchone()[0]
            if categoria:
                cur.execute("SELECT categoria_id FROM Categoria WHERE nome ILIKE %s", (categoria,))
                result = cur.fetchone()
                if result:
                    cur.execute("UPDATE FAQ SET categoria_id = %s WHERE faq_id = %s", (result[0], faq_id))
            if links_documentos:
                for link in links_documentos.split(','):
                    link = link.strip()
                    if link:
                        cur.execute(
                            "INSERT INTO FAQ_Documento (faq_id, link) VALUES (%s, %s)",
                            (faq_id, link)
                        )
        conn.commit()
        build_faiss_index()
        return jsonify({"success": True, "message": "FAQ e links inseridos com sucesso."})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        cur.close()
        conn.close()

@app.route("/upload-faq-docx-multiplos", methods=["POST"])
def upload_faq_docx_multiplos():
    if not DOCX_AVAILABLE:
        return jsonify({"success": False, "error": "python-docx não está instalado."}), 500
    
    conn = get_conn()
    cur = conn.cursor()
    if 'files' not in request.files:
        return jsonify({"success": False, "error": "Ficheiros não enviados."}), 400
    chatbot_id_raw = request.form.get("chatbot_id")
    if not chatbot_id_raw:
        return jsonify({"success": False, "error": "Chatbot ID não fornecido."}), 400
    files = request.files.getlist('files')
    total_inseridas = 0
    erros = []
    for file in files:
        try:
            doc = docx.Document(file)
            dados = {}
            for table in doc.tables:
                for row in table.rows:
                    if len(row.cells) >= 2:
                        chave_raw = row.cells[0].text.strip().lower().replace("\u2019", "'")
                        valor = row.cells[1].text.strip()
                        chave = chave_raw.replace(":", "").strip()
                        if chave and valor:
                            dados[chave] = valor
            if not dados.get("designação da faq") or not dados.get("questão") or not dados.get("resposta"):
                raise Exception("Faltam campos obrigatórios: designação, questão ou resposta.")
            designacao = dados.get("designação da faq")
            pergunta = dados.get("questão")
            resposta = dados.get("resposta")
            categoria = dados.get("categoria")
            idioma_lido = dados.get("idioma", "Português")
            idioma = normalizar_idioma(idioma_lido)
            links_documentos = ""
            for key in ["documentos associados", "links de documentos"]:
                if key in dados:
                    links_documentos = dados[key]
                    break
            chatbot_ids = []
            if chatbot_id_raw == "todos":
                cur.execute("SELECT chatbot_id FROM Chatbot")
                chatbot_ids = [row[0] for row in cur.fetchall()]
            else:
                chatbot_ids = [int(chatbot_id_raw)]
            for chatbot_id in chatbot_ids:
                cur.execute("""
                    SELECT faq_id FROM FAQ
                    WHERE chatbot_id = %s AND designacao = %s AND pergunta = %s AND resposta = %s AND idioma = %s
                """, (chatbot_id, designacao, pergunta, resposta, idioma))
                if cur.fetchone():
                    continue
                cur.execute("""
                    INSERT INTO FAQ (chatbot_id, designacao, pergunta, resposta, idioma, links_documentos)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    RETURNING faq_id
                """, (chatbot_id, designacao, pergunta, resposta, idioma, links_documentos))
                faq_id = cur.fetchone()[0]
                if categoria:
                    cur.execute("SELECT categoria_id FROM Categoria WHERE nome ILIKE %s", (categoria,))
                    result = cur.fetchone()
                    if result:
                        cur.execute("UPDATE FAQ SET categoria_id = %s WHERE faq_id = %s", (result[0], faq_id))
                if links_documentos:
                    for link in links_documentos.split(','):
                        link = link.strip()
                        if link:
                            cur.execute(
                                "INSERT INTO FAQ_Documento (faq_id, link) VALUES (%s, %s)",
                                (faq_id, link)
                            )
                total_inseridas += 1
        except Exception as e:
            erros.append(str(e))
            conn.rollback()
    conn.commit()
    build_faiss_index()
    return jsonify({"success": True, "inseridas": total_inseridas, "erros": erros})

