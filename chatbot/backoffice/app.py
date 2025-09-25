from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
import psycopg2
import docx
import PyPDF2
import logging
from fuzzywuzzy import fuzz
import faiss
from sentence_transformers import SentenceTransformer
import numpy as np
import pickle
import os
import traceback
import re
import requests
from unidecode import unidecode
from werkzeug.utils import secure_filename

# ---------- LOGGING ----------
logging.basicConfig(level=logging.DEBUG)
app = Flask(__name__)
CORS(app)

@app.route("/")
def index():
    return render_template("recursos.html")

@app.route("/contexto")
def contexto():
    return render_template("contexto.html")

@app.route("/landing")
def landing():
    return render_template("landing.html")

@app.route("/landing-2")
def landing_2():
    return render_template("landing-2.html")

@app.route("/respostas")
def respostas():
    return render_template("respostas.html")

# ---------- CONEXÃO À BASE DE DADOS ----------
try:
    conn = psycopg2.connect(
        host="localhost",
        port="5433",
        dbname="AI4Governance",
        user="postgres",
        password="29344"
    )
except Exception as e:
    print("❌ Erro ao conectar à base de dados:", e)
    raise

INDEX_PATH = "faiss.index"
FAQ_EMBEDDINGS_PATH = "faq_embeddings.pkl"
PDF_STORAGE_PATH = "pdfs/"
ICON_STORAGE_PATH = "static/icons/"
embedding_model = SentenceTransformer('all-MiniLM-L12-v2')
if not os.path.exists(PDF_STORAGE_PATH):
    os.makedirs(PDF_STORAGE_PATH)
if not os.path.exists(ICON_STORAGE_PATH):
    os.makedirs(ICON_STORAGE_PATH)

# ---------- EMBEDDINGS E FAISS ----------
SAUDACOES_PT = ["olá", "ola", "oi", "bom dia", "boa tarde", "boa noite"]
SAUDACOES_EN = ["hello", "hi", "good morning", "good afternoon", "good evening", "good night"]
SAUDACOES_PERGUNTAS_PT = ["tudo bem", "como estás", "como está", "está tudo bem", "como vais", "tá bem"]
SAUDACOES_PERGUNTAS_EN = ["how are you", "how's it going", "everything ok", "are you ok", "how are you doing"]
RESPOSTA_SAUDACAO_PT = "Olá! 👋 Como posso ajudar? Se tiver alguma dúvida, basta perguntar!"
RESPOSTA_SAUDACAO_EN = "Hello! 👋 How can I help you? If you have any question, just ask!"
RESPOSTA_TUDO_BEM_PT = "Estou sempre pronto a ajudar! 😊 Em que posso ser útil?"
RESPOSTA_TUDO_BEM_EN = "I'm always ready to help! 😊 How can I assist you?"
NEGATIVE_FEEDBACK = ["não", "nao", "não está certo", "nao esta certo", "errado", "não é isso", "nao e isso"]

def detectar_saudacao(pergunta):
    texto = pergunta.strip().lower()
    palavras = set(texto.replace("?", "").replace("!", "").replace(".", "").split())
    for saud in SAUDACOES_PT:
        for palavra in palavras:
            if saud == palavra:
                return RESPOSTA_SAUDACAO_PT
    for p in SAUDACOES_PERGUNTAS_PT:
        if p == texto:
            return RESPOSTA_TUDO_BEM_PT
    for saud in SAUDACOES_EN:
        if saud in texto:
            return RESPOSTA_SAUDACAO_EN
    for p in SAUDACOES_PERGUNTAS_EN:
        if p in texto:
            return RESPOSTA_TUDO_BEM_EN
    return None

def detectar_feedback_negativo(pergunta):
    texto = preprocess_text(pergunta.strip().lower())
    for feedback in NEGATIVE_FEEDBACK:
        if fuzz.ratio(texto, preprocess_text(feedback)) > 80:
            return True
    return False

def preprocess_text(text):
    text = unidecode(text.lower())
    text = re.sub(r'[^\w\s]', '', text)
    stop_words = {'a', 'o', 'e', 'de', 'para', 'com', 'em', 'que', 'quem', 'como'}
    text = ' '.join(word for word in text.split() if word not in stop_words)
    return text

def get_faqs_from_db(chatbot_id=None):
    cur = conn.cursor()
    if chatbot_id:
        cur.execute("SELECT faq_id, pergunta, resposta, chatbot_id FROM FAQ WHERE chatbot_id = %s", (chatbot_id,))
    else:
        cur.execute("SELECT faq_id, pergunta, resposta, chatbot_id FROM FAQ")
    return cur.fetchall()

def get_pdfs_from_db(chatbot_id=None):
    cur = conn.cursor()
    try:
        if chatbot_id:
            cur.execute("SELECT pdf_id, file_path FROM PDF_Documents WHERE chatbot_id = %s", (chatbot_id,))
        else:
            cur.execute("SELECT pdf_id, file_path FROM PDF_Documents")
        return cur.fetchall()
    finally:
        cur.close()

def build_faiss_index(chatbot_id=None):
    cur = conn.cursor()
    try:
        if chatbot_id:
            cur.execute("SELECT faq_id, pergunta, resposta, chatbot_id FROM FAQ WHERE chatbot_id = %s", (chatbot_id,))
        else:
            cur.execute("SELECT faq_id, pergunta, resposta, chatbot_id FROM FAQ")
        faqs = cur.fetchall()
        perguntas = [f[1] for f in faqs]
        if not perguntas:
            emb_dim = embedding_model.get_sentence_embedding_dimension()
            embeddings = np.zeros((1, emb_dim), dtype=np.float32)
            index = faiss.IndexFlatIP(emb_dim)
        else:
            embeddings = embedding_model.encode(perguntas, show_progress_bar=True)
            embeddings = embeddings / np.linalg.norm(embeddings, axis=1, keepdims=True)
            index = faiss.IndexFlatIP(embeddings.shape[1])
            index.add(np.array(embeddings, dtype=np.float32))
        with open(FAQ_EMBEDDINGS_PATH, 'wb') as f:
            pickle.dump({'faqs': faqs, 'embeddings': embeddings}, f)
        faiss.write_index(index, INDEX_PATH)
        logging.info(f"Índice FAISS para FAQs salvo em {INDEX_PATH}")
    except Exception as e:
        logging.error(f"Erro ao construir índice FAISS para FAQs: {e}")
        raise
    finally:
        cur.close()

def load_faiss_index():
    try:
        if not os.path.exists(INDEX_PATH) or not os.path.exists(FAQ_EMBEDDINGS_PATH):
            logging.info("Índice FAISS ou embeddings de FAQs não encontrados. Reconstruindo...")
            build_faiss_index()
        index = faiss.read_index(INDEX_PATH)
        with open(FAQ_EMBEDDINGS_PATH, 'rb') as f:
            data = pickle.load(f)
        return index, data['faqs'], data['embeddings']
    except Exception as e:
        logging.error(f"Erro ao carregar índice FAISS para FAQs: {e}")
        logging.info("Reconstruindo índice FAISS para FAQs...")
        build_faiss_index()
        index = faiss.read_index(INDEX_PATH)
        with open(FAQ_EMBEDDINGS_PATH, 'rb') as f:
            data = pickle.load(f)
        return index, data['faqs'], data['embeddings']

faiss_index, faqs_db, faq_embeddings = load_faiss_index()

def pesquisar_faiss(pergunta, chatbot_id=None, k=1, min_sim=0.7):
    pergunta = preprocess_text(pergunta)
    results = []
    if chatbot_id:
        faqs = [f for f in faqs_db if f[3] == int(chatbot_id)]
        if not faqs:
            return []
        perguntas = [preprocess_text(f[1]) for f in faqs]
        embeddings = embedding_model.encode(perguntas)
        embeddings = embeddings / np.linalg.norm(embeddings, axis=1, keepdims=True)
        index = faiss.IndexFlatIP(embeddings.shape[1])
        index.add(np.array(embeddings, dtype=np.float32))
        query_emb = embedding_model.encode([pergunta])
        query_emb = query_emb / np.linalg.norm(query_emb, axis=1, keepdims=True)
        D, I = index.search(np.array(query_emb, dtype=np.float32), min(k, len(faqs)))
        for score, idx_faq in zip(D[0], I[0]):
            if idx_faq == -1 or score < min_sim:
                continue
            faq_id, pergunta_faq, resposta_faq, chatbot_id_faq = faqs[idx_faq]
            results.append({
                'faq_id': faq_id,
                'pergunta': pergunta_faq,
                'resposta': resposta_faq,
                'score': float(score)
            })
        return results
    else:
        if len(faqs_db) == 0:
            return []
        query_emb = embedding_model.encode([pergunta])
        query_emb = query_emb / np.linalg.norm(query_emb, axis=1, keepdims=True)
        D, I = faiss_index.search(np.array(query_emb, dtype=np.float32), min(k, len(faqs_db)))
        for score, idx_faq in zip(D[0], I[0]):
            if idx_faq == -1 or score < min_sim:
                continue
            faq_id, pergunta_faq, resposta_faq, chatbot_id_faq = faqs_db[idx_faq]
            results.append({
                'faq_id': faq_id,
                'pergunta': pergunta_faq,
                'resposta': resposta_faq,
                'score': float(score)
            })
        return results

def pesquisar_pdf_ollama(pergunta, chatbot_id=None, modelo="llama3", max_chars=18000):
    pdfs = get_pdfs_from_db(chatbot_id)
    contexto_pdf = ""
    for pdf_id, file_path in pdfs:
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    texto = page.extract_text()
                    if texto:
                        contexto_pdf += texto + "\n\n"
        except Exception as e:
            print(f"Erro lendo PDF: {file_path}", e)
    contexto_pdf = contexto_pdf[:max_chars]
    prompt = f"""Responda à pergunta abaixo UTILIZANDO APENAS o conteúdo extraído dos documentos PDF. Não invente informação. Seja objetivo.
Pergunta: {pergunta}
Conteúdo dos documentos:
{contexto_pdf}
"""
    url = "http://localhost:11434/api/generate"
    payload = {
        "model": modelo,
        "prompt": prompt,
        "stream": False
    }
    try:
        resp = requests.post(url, json=payload, timeout=120)
        resp.raise_for_status()
        resposta = resp.json().get("response", "").strip()
        return resposta
    except Exception as e:
        print("Erro Ollama:", e)
        return None

def preprocess_text_for_matching(text):
    stop_words = {'a', 'o', 'e', 'de', 'para', 'com', 'em', 'que', 'quem', 'como', 'do', 'da', 'os', 'as', 'dos', 'das'}
    text = unidecode(text.lower())
    text = re.sub(r'[^\w\s]', '', text)
    text = ' '.join(word for word in text.split() if word not in stop_words)
    return text.strip()

def obter_faq_mais_semelhante(pergunta_utilizador, chatbot_id, threshold=None):
    cur = conn.cursor()
    cur.execute("SELECT pergunta, resposta FROM FAQ WHERE chatbot_id = %s", (chatbot_id,))
    faqs = cur.fetchall()
    pergunta_norm = preprocess_text_for_matching(pergunta_utilizador)
    melhor_pergunta = None
    melhor_resposta = None
    maior_score = 0
    if threshold is None:
        threshold = 40 if len(pergunta_utilizador.split()) <= 5 else 60
    for pergunta, resposta in faqs:
        pergunta_bd_norm = preprocess_text_for_matching(pergunta)
        if pergunta_norm == pergunta_bd_norm:
            return {"pergunta": pergunta, "resposta": resposta, "score": 100}
        score = fuzz.token_set_ratio(pergunta_norm, pergunta_bd_norm)
        if score > maior_score:
            maior_score = score
            melhor_pergunta = pergunta
            melhor_resposta = resposta
    if maior_score >= threshold:
        return {"pergunta": melhor_pergunta, "resposta": melhor_resposta, "score": maior_score}
    else:
        return None

def obter_mensagem_sem_resposta(chatbot_id):
    cur = conn.cursor()
    try:
        cur.execute("SELECT mensagem_sem_resposta FROM Chatbot WHERE chatbot_id = %s", (chatbot_id,))
        row = cur.fetchone()
        if row and row[0] and row[0].strip():
            return row[0].strip()
    except Exception:
        pass
    finally:
        cur.close()
    return "Não conseguimos perceber a sua questão, pode ser mais claro?"

@app.route("/chatbots/<int:chatbot_id>", methods=["DELETE"])
def eliminar_chatbot(chatbot_id):
    cur = conn.cursor()
    try:
        cur.execute("DELETE FROM FAQ_Relacionadas WHERE faq_id IN (SELECT faq_id FROM FAQ WHERE chatbot_id = %s)", (chatbot_id,))
        cur.execute("DELETE FROM FAQ_Documento WHERE faq_id IN (SELECT faq_id FROM FAQ WHERE chatbot_id = %s)", (chatbot_id,))
        cur.execute("DELETE FROM FAQ WHERE chatbot_id = %s", (chatbot_id,))
        cur.execute("DELETE FROM FonteResposta WHERE chatbot_id = %s", (chatbot_id,))
        cur.execute("DELETE FROM PDF_Documents WHERE chatbot_id = %s", (chatbot_id,))
        cur.execute("DELETE FROM Chatbot WHERE chatbot_id = %s", (chatbot_id,))
        conn.commit()
        build_faiss_index()
        global faiss_index, faqs_db, faq_embeddings
        faiss_index, faqs_db, faq_embeddings = load_faiss_index()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/categorias", methods=["GET"])
def get_categorias():
    cur = conn.cursor()
    try:
        cur.execute("SELECT categoria_id, nome FROM Categoria")
        data = cur.fetchall()
        return jsonify([{"categoria_id": c[0], "nome": c[1]} for c in data])
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/categorias", methods=["POST"])
def criar_categoria():
    cur = conn.cursor()
    data = request.get_json()
    nome = data.get("nome", "").strip()
    if not nome:
        return jsonify({"success": False, "error": "Nome da categoria é obrigatório."}), 400
    try:
        cur.execute("SELECT categoria_id FROM Categoria WHERE LOWER(nome) = LOWER(%s)", (nome,))
        if cur.fetchone():
            return jsonify({"success": False, "error": "Já existe uma categoria com esse nome."}), 409
        cur.execute(
            "INSERT INTO Categoria (nome) VALUES (%s) RETURNING categoria_id",
            (nome,)
        )
        categoria_id = cur.fetchone()[0]
        conn.commit()
        return jsonify({"categoria_id": categoria_id, "nome": nome}), 201
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        cur.close()

@app.route("/chatbots", methods=["GET"])
def get_chatbots():
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT c.chatbot_id, c.nome, c.descricao, c.data_criacao, c.cor, c.icon_path, fr.fonte,
                   array_remove(array_agg(cc.categoria_id), NULL) as categorias,
                   c.mensagem_sem_resposta
            FROM Chatbot c
            LEFT JOIN FonteResposta fr ON fr.chatbot_id = c.chatbot_id
            LEFT JOIN ChatbotCategoria cc ON cc.chatbot_id = c.chatbot_id
            GROUP BY c.chatbot_id, c.nome, c.descricao, c.data_criacao, c.cor, c.icon_path, fr.fonte, c.mensagem_sem_resposta
            ORDER BY c.chatbot_id ASC
        """)
        data = cur.fetchall()
        return jsonify([
            {
                "chatbot_id": row[0],
                "nome": row[1],
                "descricao": row[2],
                "data_criacao": row[3],
                "cor": row[4] if row[4] else "#d4af37",
                "icon_path": row[5] if row[5] else "/static/images/chatbot-icon.png",
                "fonte": row[6] if row[6] else "faq",
                "categorias": row[7] if row[7] is not None else [],
                "mensagem_sem_resposta": row[8] if len(row) > 8 else ""
            }
            for row in data
        ])
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        cur.close()

@app.route("/chatbots", methods=["POST"])
def criar_chatbot():
    cur = conn.cursor()
    data = request.get_json()
    nome = data.get("nome", "").strip()
    descricao = data.get("descricao", "").strip()
    categorias = data.get("categorias", [])
    cor = data.get("cor", "").strip() or "#d4af37"
    icon_path = data.get("icon_path", "/static/images/chatbot-icon.png")
    mensagem_sem_resposta = data.get("mensagem_sem_resposta", "").strip()
    if not nome:
        return jsonify({"success": False, "error": "Nome obrigatório."}), 400
    try:
        cur.execute("SELECT chatbot_id FROM Chatbot WHERE LOWER(nome) = LOWER(%s)", (nome,))
        if cur.fetchone():
            return jsonify({"success": False, "error": "Já existe um chatbot com esse nome."}), 409
        cur.execute(
            "INSERT INTO Chatbot (nome, descricao, cor, icon_path, mensagem_sem_resposta) VALUES (%s, %s, %s, %s, %s) RETURNING chatbot_id",
            (nome, descricao, cor, icon_path, mensagem_sem_resposta)
        )
        chatbot_id = cur.fetchone()[0]
        for categoria_id in categorias:
            cur.execute(
                "INSERT INTO ChatbotCategoria (chatbot_id, categoria_id) VALUES (%s, %s)",
                (chatbot_id, categoria_id)
            )
        cur.execute(
            "INSERT INTO FonteResposta (chatbot_id, fonte) VALUES (%s, %s)",
            (chatbot_id, "faq")
        )
        conn.commit()
        return jsonify({"success": True, "chatbot_id": chatbot_id})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        cur.close()

@app.route("/chatbots/<int:chatbot_id>", methods=["PUT"])
def atualizar_chatbot(chatbot_id):
    cur = conn.cursor()
    try:
        print("Dados recebidos:", dict(request.form))
        nome = request.form.get("nome", "").strip()
        descricao = request.form.get("descricao", "").strip()
        fonte = request.form.get("fonte", "faq")  # Valor padrão "faq" se não enviado
        categorias = request.form.getlist("categorias[]") if request.form.getlist("categorias[]") else []
        cor = request.form.get("cor", "").strip() or "#d4af37"
        mensagem_sem_resposta = request.form.get("mensagem_sem_resposta", "").strip()
        icon_path = None
        if 'icon' in request.files:
            file = request.files['icon']
            if file.filename:
                filename = secure_filename(file.filename)
                if filename:
                    icon_path = os.path.join(ICON_STORAGE_PATH, filename)
                    file.save(icon_path)
        if not nome:
            return jsonify({"success": False, "error": "O nome do chatbot é obrigatório."}), 400
        cur.execute(
            "UPDATE Chatbot SET nome=%s, descricao=%s, cor=%s, mensagem_sem_resposta=%s" + (", icon_path=%s" if icon_path else "") + " WHERE chatbot_id=%s",
            (nome, descricao, cor, mensagem_sem_resposta, icon_path, chatbot_id) if icon_path else (nome, descricao, cor, mensagem_sem_resposta, chatbot_id)
        )
        cur.execute("DELETE FROM ChatbotCategoria WHERE chatbot_id=%s", (chatbot_id,))
        for categoria_id in categorias:
            cur.execute(
                "INSERT INTO ChatbotCategoria (chatbot_id, categoria_id) VALUES (%s, %s)",
                (chatbot_id, int(categoria_id))
            )
        # Garantir que a fonte seja sempre atualizada ou inserida
        cur.execute("SELECT 1 FROM FonteResposta WHERE chatbot_id=%s", (chatbot_id,))
        if cur.fetchone():
            cur.execute("UPDATE FonteResposta SET fonte=%s WHERE chatbot_id=%s", (fonte, chatbot_id))
        else:
            cur.execute("INSERT INTO FonteResposta (chatbot_id, fonte) VALUES (%s, %s)", (chatbot_id, fonte))
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        print(traceback.format_exc())
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        cur.close()

@app.route("/chatbots/<int:chatbot_id>/categorias", methods=["GET"])
def get_categorias_chatbot(chatbot_id):
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT c.categoria_id, c.nome
            FROM Categoria c
            JOIN ChatbotCategoria cc ON c.categoria_id = cc.categoria_id
            WHERE cc.chatbot_id = %s
        """, (chatbot_id,))
        data = cur.fetchall()
        return jsonify([{"categoria_id": c[0], "nome": c[1]} for c in data])
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/chatbots/<int:chatbot_id>/categorias/<int:categoria_id>", methods=["DELETE"])
def remove_categoria_from_chatbot(chatbot_id, categoria_id):
    cur = conn.cursor()
    try:
        cur.execute("DELETE FROM ChatbotCategoria WHERE chatbot_id = %s AND categoria_id = %s", (chatbot_id, categoria_id))
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/chatbots/<int:chatbot_id>/categorias", methods=["POST"])
def add_categoria_to_chatbot(chatbot_id):
    cur = conn.cursor()
    data = request.get_json()
    categoria_id = data.get("categoria_id")
    if not categoria_id:
        return jsonify({"success": False, "error": "ID da categoria é obrigatório."}), 400
    try:
        cur.execute("INSERT INTO ChatbotCategoria (chatbot_id, categoria_id) VALUES (%s, %s) ON CONFLICT DO NOTHING", (chatbot_id, categoria_id))
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/fonte/<int:chatbot_id>", methods=["GET"])
def obter_fonte_chatbot(chatbot_id):
    cur = conn.cursor()
    try:
        cur.execute("SELECT fonte FROM FonteResposta WHERE chatbot_id = %s", (chatbot_id,))
        row = cur.fetchone()
        if row:
            fonte = row[0] if row[0] else "faq"
            return jsonify({"success": True, "fonte": fonte})
        return jsonify({"success": False, "erro": "Chatbot não encontrado."}), 404
    except Exception as e:
        return jsonify({"success": False, "erro": str(e)}), 500

@app.route("/fonte", methods=["POST"])
def definir_fonte_chatbot():
    cur = conn.cursor()
    data = request.get_json()
    chatbot_id = data.get("chatbot_id")
    fonte = data.get("fonte")
    if fonte not in ["faq", "faiss", "faq+raga"]:
        return jsonify({"success": False, "erro": "Fonte inválida."}), 400
    try:
        cur.execute("SELECT 1 FROM FonteResposta WHERE chatbot_id = %s", (chatbot_id,))
        if not cur.fetchone():
            cur.execute("INSERT INTO FonteResposta (chatbot_id, fonte) VALUES (%s, %s)", (chatbot_id, fonte))
        else:
            cur.execute("UPDATE FonteResposta SET fonte = %s WHERE chatbot_id = %s", (fonte, chatbot_id))
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "erro": str(e)}), 500

@app.route("/faq-categoria/<categoria>", methods=["GET"])
def obter_faq_por_categoria(categoria):
    cur = conn.cursor()
    try:
        chatbot_id = request.args.get("chatbot_id")
        if not chatbot_id:
            return jsonify({"success": False, "erro": "chatbot_id não fornecido."}), 400
        cur.execute("""
            SELECT f.faq_id, f.pergunta, f.resposta
            FROM FAQ f
            INNER JOIN Categoria c ON f.categoria_id = c.categoria_id
            WHERE LOWER(c.nome) = LOWER(%s) AND f.chatbot_id = %s
            ORDER BY RANDOM()
            LIMIT 1
        """, (categoria, chatbot_id))
        resultado = cur.fetchone()
        if resultado:
            return jsonify({
                "success": True,
                "faq_id": resultado[0],
                "pergunta": resultado[1],
                "resposta": resultado[2]
            })
        else:
            return jsonify({
                "success": False,
                "erro": f"Nenhuma FAQ encontrada para a categoria '{categoria}'."
            }), 404
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "erro": str(e)}), 500

@app.route("/faqs", methods=["GET"])
def get_faqs():
    cur = conn.cursor()
    try:
        cur.execute("SELECT faq_id, chatbot_id, designacao, pergunta, resposta FROM FAQ")
        data = cur.fetchall()
        return jsonify([
            {"faq_id": f[0], "chatbot_id": f[1], "designacao": f[2], "pergunta": f[3], "resposta": f[4]} for f in data
        ])
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/faqs/<int:faq_id>", methods=["GET"])
def get_faq_by_id(faq_id):
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT f.faq_id, f.chatbot_id, f.categoria_id, f.designacao, f.pergunta, f.resposta, f.idioma, f.links_documentos,
                   c.nome as categoria_nome, f.recomendado
            FROM FAQ f
            LEFT JOIN Categoria c ON f.categoria_id = c.categoria_id
            WHERE f.faq_id = %s
        """, (faq_id,))
        faq = cur.fetchone()
        if not faq:
            return jsonify({"success": False, "error": "FAQ não encontrada."}), 404
        return jsonify({
            "success": True,
            "faq": {
                "faq_id": faq[0],
                "chatbot_id": faq[1],
                "categoria_id": faq[2],
                "designacao": faq[3],
                "pergunta": faq[4],
                "resposta": faq[5],
                "idioma": faq[6],
                "links_documentos": faq[7],
                "categoria_nome": faq[8],
                "recomendado": faq[9]
            }
        })
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/faqs/chatbot/<int:chatbot_id>", methods=["GET"])
def get_faqs_por_chatbot(chatbot_id):
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT f.faq_id, c.nome, f.pergunta, f.resposta
            FROM FAQ f
            LEFT JOIN Categoria c ON f.categoria_id = c.categoria_id
            WHERE f.chatbot_id = %s
        """, (chatbot_id,))
        data = cur.fetchall()
        return jsonify([
            {
                "faq_id": row[0],
                "categoria": row[1],
                "pergunta": row[2],
                "resposta": row[3]
            }
            for row in data
        ])
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/faqs/<int:faq_id>", methods=["PUT"])
def update_faq(faq_id):
    cur = conn.cursor()
    data = request.get_json()
    try:
        pergunta = data.get("pergunta", "").strip()
        resposta = data.get("resposta", "").strip()
        idioma = data.get("idioma", "pt").strip()
        categorias = data.get("categorias", [])
        recomendado = data.get("recomendado", False)
        categoria_id = categorias[0] if categorias else None
        cur.execute("""
            UPDATE FAQ SET pergunta=%s, resposta=%s, idioma=%s, categoria_id=%s, recomendado=%s
            WHERE faq_id=%s
        """, (pergunta, resposta, idioma, categoria_id, recomendado, faq_id))
        conn.commit()
        build_faiss_index()
        global faiss_index, faqs_db, faq_embeddings
        faiss_index, faqs_db, faq_embeddings = load_faiss_index()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/faqs", methods=["POST"])
def add_faq():
    global faiss_index, faqs_db, faq_embeddings
    cur = conn.cursor()
    data = request.get_json()
    try:
        idioma = data.get("idioma", "").strip()
        if not idioma:
            return jsonify({"success": False, "error": "O campo 'idioma' é obrigatório."}), 400
        links_documentos = data.get("links_documentos", "").strip()
        recomendado = data.get("recomendado", False)
        cur.execute("""
            SELECT faq_id FROM FAQ
            WHERE chatbot_id = %s AND designacao = %s AND pergunta = %s AND resposta = %s AND idioma = %s
        """, (data["chatbot_id"], data["designacao"], data["pergunta"], data["resposta"], idioma))
        if cur.fetchone():
            return jsonify({"success": False, "error": "Esta FAQ já está inserida."}), 409
        cur.execute("""
            INSERT INTO FAQ (chatbot_id, categoria_id, designacao, pergunta, resposta, idioma, links_documentos, recomendado)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING faq_id
        """, (
            data["chatbot_id"],
            data["categoria_id"],
            data["designacao"],
            data["pergunta"],
            data["resposta"],
            idioma,
            links_documentos,
            recomendado
        ))
        faq_id = cur.fetchone()[0]
        if links_documentos:
            for link in links_documentos.split(','):
                link = link.strip()
                if link:
                    cur.execute(
                        "INSERT INTO FAQ_Documento (faq_id, link) VALUES (%s, %s)",
                        (faq_id, link)
                    )
        if "relacionadas" in data and data["relacionadas"].strip():
            for rel_id in data["relacionadas"].split(','):
                cur.execute("INSERT INTO FAQ_Relacionadas (faq_id, faq_relacionada_id) VALUES (%s, %s)", (faq_id, int(rel_id.strip())))
        conn.commit()
        build_faiss_index()
        faiss_index, faqs_db, faq_embeddings = load_faiss_index()
        return jsonify({"success": True, "faq_id": faq_id})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/faqs/<int:faq_id>", methods=["DELETE"])
def delete_faq(faq_id):
    cur = conn.cursor()
    try:
        cur.execute("DELETE FROM FAQ WHERE faq_id = %s", (faq_id,))
        conn.commit()
        build_faiss_index()
        global faiss_index, faqs_db, faq_embeddings
        faiss_index, faqs_db, faq_embeddings = load_faiss_index()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/upload-pdf", methods=["POST"])
def upload_pdf():
    cur = conn.cursor()
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "Ficheiro não enviado."}), 400
    files = request.files.getlist('file')
    chatbot_id = request.form.get("chatbot_id")
    if not chatbot_id:
        return jsonify({"success": False, "error": "Chatbot ID não fornecido."}), 400
    uploaded_pdf_ids = []
    try:
        for file in files:
            filename = file.filename
            if not filename.lower().endswith('.pdf'):
                return jsonify({"success": False, "error": "Apenas ficheiros PDF são permitidos."}), 400
            file_path = os.path.join(PDF_STORAGE_PATH, filename)
            file.save(file_path)
            file.close()
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                if reader.is_encrypted:
                    return jsonify({"success": False, "error": f"O PDF '{filename}' está protegido por senha."}), 400
                if not reader.pages:
                    return jsonify({"success": False, "error": f"O PDF '{filename}' está vazio ou corrompido."}), 400
            cur.execute(
                "INSERT INTO PDF_Documents (chatbot_id, filename, file_path) VALUES (%s, %s, %s) RETURNING pdf_id",
                (chatbot_id, filename, file_path)
            )
            pdf_id = cur.fetchone()[0]
            uploaded_pdf_ids.append(pdf_id)
        conn.commit()
        return jsonify({"success": True, "pdf_ids": uploaded_pdf_ids, "message": "PDF(s) carregado(s) com sucesso."})
    except PyPDF2.errors.PdfReadError:
        return jsonify({"success": False, "error": "Erro ao ler o PDF. Verifique se o arquivo não está corrompido."}), 400
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/upload-faq-docx", methods=["POST"])
def upload_faq_docx():
    cur = conn.cursor()
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "Ficheiro não enviado."}), 400
    file = request.files['file']
    chatbot_id_raw = request.form.get("chatbot_id")
    if not chatbot_id_raw:
        return jsonify({"success": False, "error": "Chatbot ID não fornecido."}), 400
    def normalizar_idioma(valor):
        if not valor:
            return "pt"
        valor = valor.strip().lower()
        if valor.startswith("port"):
            return "pt"
        if valor.startswith("ingl"):
            return "en"
        return valor[:2]
    try:
        doc = docx.Document(file)
        dados = {}
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    chave_raw = row.cells[0].text.strip().lower().replace("’", "'")
                    valor = row.cells[1].text.strip()
                    chave = chave_raw.replace(":", "").strip()
                    if chave and valor:
                        dados[chave] = valor
        if not dados.get("designação da faq") or not dados.get("questão") or not dados.get("resposta"):
            raise Exception("Faltam campos obrigatórios: designação, questão ou resposta.")
        designacao = dados.get("designação da faq")
        pergunta = dados.get("questão")
        resposta = dados.get("resposta")
        categoria = dados.get("categoria")
        idioma_lido = dados.get("idioma", "Português")
        idioma = normalizar_idioma(idioma_lido)
        links_documentos = ""
        for key in ["documentos associados", "links de documentos"]:
            if key in dados:
                links_documentos = dados[key]
                break
        chatbot_ids = []
        if chatbot_id_raw == "todos":
            cur.execute("SELECT chatbot_id FROM Chatbot")
            chatbot_ids = [row[0] for row in cur.fetchall()]
        else:
            chatbot_ids = [int(chatbot_id_raw)]
        for chatbot_id in chatbot_ids:
            cur.execute("""
                SELECT faq_id FROM FAQ
                WHERE chatbot_id = %s AND designacao = %s AND pergunta = %s AND resposta = %s AND idioma = %s
            """, (chatbot_id, designacao, pergunta, resposta, idioma))
            if cur.fetchone():
                continue
            cur.execute("""
                INSERT INTO FAQ (chatbot_id, designacao, pergunta, resposta, idioma, links_documentos)
                VALUES (%s, %s, %s, %s, %s, %s)
                RETURNING faq_id
            """, (chatbot_id, designacao, pergunta, resposta, idioma, links_documentos))
            faq_id = cur.fetchone()[0]
            if categoria:
                cur.execute("SELECT categoria_id FROM Categoria WHERE nome ILIKE %s", (categoria,))
                result = cur.fetchone()
                if result:
                    cur.execute("UPDATE FAQ SET categoria_id = %s WHERE faq_id = %s", (result[0], faq_id))
            if links_documentos:
                for link in links_documentos.split(','):
                    link = link.strip()
                    if link:
                        cur.execute(
                            "INSERT INTO FAQ_Documento (faq_id, link) VALUES (%s, %s)",
                            (faq_id, link)
                        )
        conn.commit()
        build_faiss_index()
        global faiss_index, faqs_db, faq_embeddings
        faiss_index, faqs_db, faq_embeddings = load_faiss_index()
        return jsonify({"success": True, "message": "FAQ e links inseridos com sucesso."})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/upload-faq-docx-multiplos", methods=["POST"])
def upload_faq_docx_multiplos():
    cur = conn.cursor()
    if 'files' not in request.files:
        return jsonify({"success": False, "error": "Ficheiros não enviados."}), 400
    chatbot_id_raw = request.form.get("chatbot_id")
    if not chatbot_id_raw:
        return jsonify({"success": False, "error": "Chatbot ID não fornecido."}), 400
    files = request.files.getlist('files')
    total_inseridas = 0
    erros = []
    def normalizar_idioma(valor):
        if not valor:
            return "pt"
        valor = valor.strip().lower()
        if valor.startswith("port"):
            return "pt"
        if valor.startswith("ingl"):
            return "en"
        return valor[:2]
    for file in files:
        try:
            doc = docx.Document(file)
            dados = {}
            for table in doc.tables:
                for row in table.rows:
                    if len(row.cells) >= 2:
                        chave_raw = row.cells[0].text.strip().lower().replace("’", "'")
                        valor = row.cells[1].text.strip()
                        chave = chave_raw.replace(":", "").strip()
                        if chave and valor:
                            dados[chave] = valor
            if not dados.get("designação da faq") or not dados.get("questão") or not dados.get("resposta"):
                raise Exception("Faltam campos obrigatórios: designação, questão ou resposta.")
            designacao = dados.get("designação da faq")
            pergunta = dados.get("questão")
            resposta = dados.get("resposta")
            categoria = dados.get("categoria")
            idioma_lido = dados.get("idioma", "Português")
            idioma = normalizar_idioma(idioma_lido)
            links_documentos = ""
            for key in ["documentos associados", "links de documentos"]:
                if key in dados:
                    links_documentos = dados[key]
                    break
            chatbot_ids = []
            if chatbot_id_raw == "todos":
                cur.execute("SELECT chatbot_id FROM Chatbot")
                chatbot_ids = [row[0] for row in cur.fetchall()]
            else:
                chatbot_ids = [int(chatbot_id_raw)]
            for chatbot_id in chatbot_ids:
                cur.execute("""
                    SELECT faq_id FROM FAQ
                    WHERE chatbot_id = %s AND designacao = %s AND pergunta = %s AND resposta = %s AND idioma = %s
                """, (chatbot_id, designacao, pergunta, resposta, idioma))
                if cur.fetchone():
                    continue
                cur.execute("""
                    INSERT INTO FAQ (chatbot_id, designacao, pergunta, resposta, idioma, links_documentos)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    RETURNING faq_id
                """, (chatbot_id, designacao, pergunta, resposta, idioma, links_documentos))
                faq_id = cur.fetchone()[0]
                if categoria:
                    cur.execute("SELECT categoria_id FROM Categoria WHERE nome ILIKE %s", (categoria,))
                    result = cur.fetchone()
                    if result:
                        cur.execute("UPDATE FAQ SET categoria_id = %s WHERE faq_id = %s", (result[0], faq_id))
                if links_documentos:
                    for link in links_documentos.split(','):
                        link = link.strip()
                        if link:
                            cur.execute(
                                "INSERT INTO FAQ_Documento (faq_id, link) VALUES (%s, %s)",
                                (faq_id, link)
                            )
                total_inseridas += 1
        except Exception as e:
            erros.append(str(e))
            conn.rollback()
    conn.commit()
    build_faiss_index()
    global faiss_index, faqs_db, faq_embeddings
    faiss_index, faqs_db, faq_embeddings = load_faiss_index()
    return jsonify({"success": True, "inseridas": total_inseridas, "erros": erros})

@app.route("/faqs/detalhes", methods=["GET"])
def get_faqs_detalhes():
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT f.faq_id, f.chatbot_id, f.designacao, f.pergunta, f.resposta, f.idioma, f.links_documentos,
                   f.categoria_id, c.nome AS categoria_nome, ch.nome AS chatbot_nome, f.recomendado
            FROM FAQ f
            LEFT JOIN Categoria c ON f.categoria_id = c.categoria_id
            LEFT JOIN Chatbot ch ON f.chatbot_id = ch.chatbot_id
            ORDER BY f.faq_id
        """)
        data = cur.fetchall()
        return jsonify([
            {
                "faq_id": r[0],
                "chatbot_id": r[1],
                "designacao": r[2],
                "pergunta": r[3],
                "resposta": r[4],
                "idioma": r[5],
                "links_documentos": r[6],
                "categoria_id": r[7],
                "categoria_nome": r[8],
                "chatbot_nome": r[9],
                "recomendado": r[10]
            }
            for r in data
        ])
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
@app.route("/obter-resposta", methods=["POST"])
def obter_resposta():
    cur = conn.cursor()
    dados = request.get_json()
    pergunta = dados.get("pergunta", "").strip()
    chatbot_id = dados.get("chatbot_id")
    fonte = dados.get("fonte", "faq")
    idioma = dados.get("idioma", "pt")
    feedback = dados.get("feedback", None)
    print("DEBUG /obter-resposta:", {
        "pergunta": pergunta,
        "chatbot_id": chatbot_id,
        "fonte": fonte,
        "feedback": feedback,
        "type_feedback": type(feedback)
    })
    try:
        chatbot_id = int(chatbot_id)
    except Exception:
        return jsonify({"success": False, "erro": "Chatbot ID inválido."}), 400
    saudacao = detectar_saudacao(pergunta)
    if saudacao:
        return jsonify({
            "success": True,
            "fonte": "SAUDACAO",
            "resposta": saudacao,
            "faq_id": None,
            "categoria_id": None,
            "pergunta_faq": None,
            "documentos": []
        })
    if not pergunta or (len(pergunta) < 4 and not any(char.isalpha() for char in pergunta)):
        return jsonify({
            "success": False,
            "erro": "Pergunta demasiado curta ou não reconhecida como válida."
        })
    if fonte == "faq+raga" and (feedback is None or feedback == "") and pergunta.lower() in ["sim", "yes"]:
        return jsonify({
            "success": False,
            "erro": "Por favor utilize os botões abaixo para confirmar.",
            "prompt_rag": True
        })
    try:
        if fonte == "faq":
            resultado = obter_faq_mais_semelhante(pergunta, chatbot_id)
            if resultado:
                cur.execute("""
                    SELECT faq_id, categoria_id FROM FAQ
                    WHERE LOWER(pergunta) = LOWER(%s) AND chatbot_id = %s
                """, (resultado["pergunta"], chatbot_id))
                row = cur.fetchone()
                faq_id, categoria_id = row if row else (None, None)
                cur.execute("SELECT link FROM FAQ_Documento WHERE faq_id = %s", (faq_id,))
                docs = [r[0] for r in cur.fetchall()]
                return jsonify({
                    "success": True,
                    "fonte": "FAQ",
                    "resposta": resultado["resposta"],
                    "faq_id": faq_id,
                    "categoria_id": categoria_id,
                    "score": resultado["score"],
                    "pergunta_faq": resultado["pergunta"],
                    "documentos": docs
                })
            return jsonify({
                "success": False,
                "erro": obter_mensagem_sem_resposta(chatbot_id)
            })
        elif fonte == "faiss":
            faiss_resultados = pesquisar_faiss(pergunta, chatbot_id=chatbot_id, k=1, min_sim=0.7)
            if faiss_resultados:
                faq_id = faiss_resultados[0]['faq_id']
                cur.execute("SELECT link FROM FAQ_Documento WHERE faq_id = %s", (faq_id,))
                docs = [r[0] for r in cur.fetchall()]
                return jsonify({
                    "success": True,
                    "fonte": "FAISS",
                    "resposta": faiss_resultados[0]['resposta'],
                    "faq_id": faq_id,
                    "score": faiss_resultados[0]['score'],
                    "pergunta_faq": faiss_resultados[0]['pergunta'],
                    "documentos": docs
                })
            else:
                resultado = obter_faq_mais_semelhante(pergunta, chatbot_id, threshold=80)
                if resultado:
                    cur.execute("""
                        SELECT faq_id, categoria_id FROM FAQ
                        WHERE LOWER(pergunta) = LOWER(%s) AND chatbot_id = %s
                    """, (resultado["pergunta"], chatbot_id))
                    row = cur.fetchone()
                    faq_id, categoria_id = row if row else (None, None)
                    cur.execute("SELECT link FROM FAQ_Documento WHERE faq_id = %s", (faq_id,))
                    docs = [r[0] for r in cur.fetchall()]
                    return jsonify({
                        "success": True,
                        "fonte": "FUZZY",
                        "resposta": resultado["resposta"],
                        "faq_id": faq_id,
                        "categoria_id": categoria_id,
                        "score": resultado["score"],
                        "pergunta_faq": resultado["pergunta"],
                        "documentos": docs
                    })
                return jsonify({
                    "success": False,
                    "erro": "Não encontrei nenhuma resposta suficientemente semelhante na base de dados."
                })
        elif fonte == "faq+raga":
            resultado = obter_faq_mais_semelhante(pergunta, chatbot_id)
            if resultado:
                cur.execute("""
                    SELECT faq_id, categoria_id FROM FAQ
                    WHERE LOWER(pergunta) = LOWER(%s) AND chatbot_id = %s
                """, (resultado["pergunta"], chatbot_id))
                row = cur.fetchone()
                faq_id, categoria_id = row if row else (None, None)
                cur.execute("SELECT link FROM FAQ_Documento WHERE faq_id = %s", (faq_id,))
                docs = [r[0] for r in cur.fetchall()]
                return jsonify({
                    "success": True,
                    "fonte": "FAQ",
                    "resposta": resultado["resposta"],
                    "faq_id": faq_id,
                    "categoria_id": categoria_id,
                    "score": resultado["score"],
                    "pergunta_faq": resultado["pergunta"],
                    "documentos": docs
                })
            elif feedback and feedback.strip().lower() == "try_rag":
                print("DEBUG: A tentar responder via RAG (PDF) via Ollama")
                resposta_ollama = pesquisar_pdf_ollama(pergunta, chatbot_id=chatbot_id)
                if resposta_ollama:
                    pdfs = get_pdfs_from_db(chatbot_id)
                    file_path = pdfs[0][1] if pdfs else None
                    return jsonify({
                        "success": True,
                        "fonte": "RAG-OLLAMA",
                        "resposta": resposta_ollama,
                        "faq_id": None,
                        "categoria_id": None,
                        "score": None,
                        "pergunta_faq": None,
                        "documentos": [file_path] if file_path else []
                    })
                else:
                    return jsonify({
                        "success": False,
                        "erro": "Não foi possível encontrar uma resposta nos documentos PDF usando Ollama."
                    })
            else:
                print("DEBUG: feedback != 'try_rag' -> devolve prompt_rag")
                return jsonify({
                    "success": False,
                    "erro": "Pergunta não encontrada nas FAQs. Deseja tentar encontrar uma resposta nos documentos PDF? Isso pode levar alguns segundos.",
                    "prompt_rag": True
                })
        else:
            return jsonify({"success": False, "erro": "Fonte inválida."}), 400
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"success": False, "erro": str(e)}), 500
    finally:
        cur.close()
@app.route("/perguntas-semelhantes", methods=["POST"])
def perguntas_semelhantes():
    cur = conn.cursor()
    dados = request.get_json()
    pergunta_atual = dados.get("pergunta", "")
    chatbot_id = dados.get("chatbot_id")
    idioma = dados.get("idioma", "pt")
    try:
        cur.execute("""
            SELECT categoria_id
            FROM FAQ
            WHERE LOWER(pergunta) = LOWER(%s) AND chatbot_id = %s AND idioma = %s
        """, (pergunta_atual.strip().lower(), chatbot_id, idioma))
        categoria_row = cur.fetchone()
        if not categoria_row or categoria_row[0] is None:
            return jsonify({"success": True, "sugestoes": []})
        categoria_id = categoria_row[0]
        cur.execute("""
            SELECT pergunta
            FROM FAQ
            WHERE categoria_id = %s
              AND recomendado = TRUE
              AND LOWER(pergunta) != LOWER(%s)
              AND chatbot_id = %s
              AND idioma = %s
            ORDER BY RANDOM()
            LIMIT 2
        """, (categoria_id, pergunta_atual.strip().lower(), chatbot_id, idioma))
        sugestoes = [row[0] for row in cur.fetchall()]
        return jsonify({"success": True, "sugestoes": sugestoes})
    except Exception as e:
        return jsonify({"success": False, "erro": str(e)}), 500
@app.route("/faqs-aleatorias", methods=["POST"])
def faqs_aleatorias():
    cur = conn.cursor()
    dados = request.get_json()
    idioma = dados.get("idioma", "pt")
    n = int(dados.get("n", 3))
    chatbot_id = dados.get("chatbot_id")
    try:
        if chatbot_id:
            cur.execute("""
                SELECT pergunta
                FROM FAQ
                WHERE idioma = %s AND chatbot_id = %s
                ORDER BY RANDOM()
                LIMIT %s
            """, (idioma, chatbot_id, n))
        else:
            cur.execute("""
                SELECT pergunta
                FROM FAQ
                WHERE idioma = %s
                ORDER BY RANDOM()
                LIMIT %s
            """, (idioma, n))
        faqs = [row[0] for row in cur.fetchall()]
        return jsonify({"success": True, "faqs": [{"pergunta": p} for p in faqs]})
    except Exception as e:
        return jsonify({"success": False, "erro": str(e)}), 500
@app.route("/chatbot/<int:chatbot_id>", methods=["GET"])
def obter_nome_chatbot(chatbot_id):
    cur = conn.cursor()
    try:
        cur.execute("SELECT nome, cor, icon_path FROM Chatbot WHERE chatbot_id = %s", (chatbot_id,))
        row = cur.fetchone()
        if row:
            return jsonify({"success": True, "nome": row[0], "cor": row[1] or "#d4af37", "icon": row[2] or "/static/images/chatbot-icon.png"})
        return jsonify({"success": False, "erro": "Chatbot não encontrado."}), 404
    except Exception as e:
        return jsonify({"success": False, "erro": str(e)}), 500
@app.route("/rebuild-faiss", methods=["POST"])
def rebuild_faiss():
    build_faiss_index()
    global faiss_index, faqs_db, faq_embeddings
    faiss_index, faqs_db, faq_embeddings = load_faiss_index()
    return jsonify({"success": True, "msg": "FAISS index rebuilt."})
# ---------- INICIAR SERVIDOR ----------
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)